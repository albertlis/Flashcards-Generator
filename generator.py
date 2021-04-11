from math import ceil
from docx import Document
from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, Length
from docx.table import Table
import argparse


def setup_margins(doc: Document, margin: float):
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(margin)
        section.bottom_margin = Cm(margin)
        section.left_margin = Cm(margin)
        section.right_margin = Cm(margin)
        section.page_height = Cm(29.7)
        section.page_width = Cm(21)


def setup_table(table: Table, cell_height: float):
    table.allow_autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        row.height = Cm(cell_height)
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def write_texts(cell, text: str, fontsize: int = 14):
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = True
    font = run.font
    font.size = Pt(fontsize)


def write_translations(cell, text: str, fontsize: int = 14):
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = True
    font = run.font
    font.size = Pt(fontsize)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description='Generate flashcards for language learning')
    parser.add_argument('-w', '--words', type=str, help='txt file with words', required=True, metavar='file.txt')
    parser.add_argument('-t', '--translations', type=str, help='txt file with translations', metavar='file.txt',
                        default=None)
    parser.add_argument('--cell_height', type=float, help='cells height in Cm', default=3.0, metavar='number')
    parser.add_argument('-m', '--margin', type=float, help='page margins in Cm', default=0.5, metavar='number')
    parser.add_argument('-f', '--filename', type=str, help='Flashcards result filename', default='flashcards.docx',
                        metavar='file.docx')
    parser.add_argument('-c', '--columns', type=int, help='Number of table columns', default=5,
                        metavar='integer_number')
    parser.add_argument('-rr', '--remove_repetitions', type=bool, help='Remove repetitions from words file',
                        default=False, metavar='boolean')

    return parser.parse_args()


def read_files(words_path: str, translations_path: str = None):
    with open(words_path, 'r', encoding='utf8') as f:
        text_lines = [line.strip().lower() for line in f.readlines()]

    translations_lines = None
    if translations_path:
        with open(translations_path, 'r', encoding='utf8') as f:
            translations_lines = [line.strip().lower() for line in f.readlines()]

    return text_lines, translations_lines


def create_tables(doc, text_lines, cell_height, margin, cols):
    nr_of_rows = remaining_rows = ceil(len(text_lines) / cols)
    page_height = (doc.sections[0].page_height / Length._EMUS_PER_CM) - (2 * margin)
    nr_of_tables = ceil(nr_of_rows * cell_height / page_height)
    page_rows = int(page_height // cell_height)
    tables = []
    translate_tables = []
    for i in range(nr_of_tables):
        if remaining_rows >= page_rows:
            rows = page_rows
        else:
            rows = remaining_rows
        remaining_rows -= rows

        table = doc.add_table(rows, cols, style='Table Grid')
        doc.add_page_break()
        translate_table = doc.add_table(rows, cols, style='Table Grid')
        if i != nr_of_tables - 1:
            doc.add_page_break()

        setup_table(table, cell_height)
        setup_table(translate_table, cell_height)

        tables.append(table)
        translate_tables.append(translate_table)
    return tables, translate_tables


def fill_tables(tables, translate_tables, text_lines, translations_lines):
    counter = 0
    for i, (table, translate_table) in enumerate(zip(tables, translate_tables)):
        for row, translate_row in zip(table.rows, translate_table.rows):
            for cell, translate_cell in zip(row.cells, reversed(translate_row.cells)):
                if counter >= len(text_lines):
                    break
                write_texts(cell, text_lines[counter])
                write_translations(translate_cell, translations_lines[counter])
                counter += 1


def remove_repetitions(text_lines):
    cleared_lines = set(text_lines)
    if len(cleared_lines) != len(text_lines):
        print(f'Removed {len(text_lines) - len(cleared_lines)} repetitions, check out your translations')
    return set(text_lines)


def main():
    args = parse_args()
    cell_height = args.cell_height
    margin = args.margin

    doc = Document()
    setup_margins(doc, margin)

    filename = args.filename
    cols = args.columns

    text_lines, translations_lines = read_files(args.words, args.translations)

    if args.remove_repetitions:
        text_lines = remove_repetitions(text_lines)
        print('Saving text without repetitions')
        with open('words_without_repetitions.txt', 'w', encoding='utf8') as f:
            f.write('\n'.join(text_lines))

    tables, translate_tables = create_tables(doc, text_lines, cell_height, margin, cols)
    fill_tables(tables, translate_tables, text_lines, translations_lines)
    doc.save(filename)


if __name__ == '__main__':
    main()
